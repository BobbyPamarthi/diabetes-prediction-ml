"""Generate a professional Word document for the Diabetes Prediction project."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '4F46E5'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def rgb(r,g,b): return RGBColor(r,g,b)

def para(text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0]
    if level == 1:
        run.font.color.rgb = rgb(30, 64, 175)   # deep blue
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = rgb(79, 70, 229)   # indigo
        run.font.size = Pt(14)
    else:
        run.font.color.rgb = rgb(100, 116, 139)
        run.font.size = Pt(12)
    return h

def bullet(text, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def hline():
    p  = doc.add_paragraph()
    pr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '4F46E5')
    pb.append(bot)
    pr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DIABETES PREDICTION USING MACHINE LEARNING')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = rgb(30, 64, 175)

para('A Comparative Study of Classification Algorithms with Ensemble Learning',
     size=13, color=rgb(100,116,139), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

hline()

para('Project Report', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('Language: Python  |  Framework: Scikit-learn, TensorFlow, Flask',
     size=11, color=rgb(100,116,139), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('Dataset: diabetes_prediction_dataset.csv  (100,000 records)',
     size=11, color=rgb(100,116,139), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual)
# ═══════════════════════════════════════════════════════════════════════════════
heading('Table of Contents', level=1)
hline()

toc_items = [
    ('1.', 'Project Overview'),
    ('2.', 'Dataset Description'),
    ('3.', 'Data Preprocessing'),
    ('4.', 'Models Used'),
    ('5.', 'Model Results & Comparison'),
    ('6.', 'ANN Architecture & Training'),
    ('7.', 'Ensemble Model'),
    ('8.', 'Web Application'),
    ('9.', 'Project File Structure'),
    ('10.','How to Run the Project'),
    ('11.','Conclusion'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True
    r1.font.color.rgb = rgb(79, 70, 229)
    r1.font.size = Pt(11)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading('1. Project Overview', level=1)
hline()
para(
    'This project develops a machine learning system to predict the risk of diabetes '
    'in patients based on clinical and demographic features. Nine classification '
    'algorithms — ranging from traditional statistical models to deep learning — are '
    'trained, evaluated, and compared. A final ensemble model combines the three best '
    'performers to achieve the highest possible predictive accuracy.',
    size=11, space_after=8
)
para('Key objectives:', bold=True, size=11, space_after=4)
bullet('Compare 9 ML models on a large, real-world diabetes dataset')
bullet('Handle class imbalance using SMOTE oversampling')
bullet('Build an interactive web dashboard to visualise model performance')
bullet('Provide a live patient-facing prediction form (Random Forest backend)')
bullet('Achieve > 97% accuracy through ensemble learning')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. DATASET
# ═══════════════════════════════════════════════════════════════════════════════
heading('2. Dataset Description', level=1)
hline()
para('File: diabetes_prediction_dataset.csv', bold=True, size=11)
para(
    'The dataset contains 100,000 patient records with 8 input features and one binary '
    'target variable (diabetes: 0 = No, 1 = Yes).  The original class distribution is '
    'heavily imbalanced: 91,500 non-diabetic vs 8,500 diabetic patients.',
    size=11, space_after=10
)

heading('2.1  Feature Descriptions', level=2)
cols = ['Feature', 'Type', 'Description', 'Example Values']
rows = [
    ['gender',             'Categorical', 'Patient gender',                    'Female, Male, Other'],
    ['age',                'Numeric',     'Age in years',                      '1.0 – 80.0'],
    ['hypertension',       'Binary',      'Has hypertension (high BP)',         '0 = No, 1 = Yes'],
    ['heart_disease',      'Binary',      'Has heart disease',                 '0 = No, 1 = Yes'],
    ['smoking_history',    'Categorical', 'Smoking status',                    'never, current, former…'],
    ['bmi',                'Numeric',     'Body Mass Index',                   '10.0 – 95.7'],
    ['HbA1c_level',        'Numeric',     'Glycated haemoglobin (%)',          '3.5 – 9.0'],
    ['blood_glucose_level','Numeric',     'Blood glucose (mg/dL)',             '80 – 300'],
    ['diabetes',           'Binary',      'Target: diabetic or not',           '0 = No, 1 = Yes'],
]
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(cols):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, '1E3A8A')
    p2 = cell.paragraphs[0]
    run = p2.add_run(h)
    run.bold = True
    run.font.color.rgb = rgb(255,255,255)
    run.font.size = Pt(10)
for i, row_data in enumerate(rows):
    row = t.add_row()
    bg = 'F0F4FF' if i % 2 == 0 else 'FFFFFF'
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        p2 = cell.paragraphs[0]
        run = p2.add_run(val)
        run.font.size = Pt(10)
        if j == 0:
            run.bold = True
            run.font.color.rgb = rgb(79,70,229)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. PREPROCESSING
# ═══════════════════════════════════════════════════════════════════════════════
heading('3. Data Preprocessing', level=1)
hline()

heading('3.1  Categorical Encoding', level=2)
para(
    'Categorical columns (gender, smoking_history) are converted to numeric using '
    'one-hot encoding via pd.get_dummies(drop_first=True). This produces 13 final '
    'input features after encoding.',
    size=11, space_after=8
)

heading('3.2  SMOTE — Handling Class Imbalance', level=2)
para(
    'The original dataset is imbalanced (91,500 : 8,500). Synthetic Minority '
    'Over-sampling Technique (SMOTE) is applied to the training data to create '
    'synthetic diabetic samples, balancing both classes to 91,500 each before training.',
    size=11, space_after=6
)
p = doc.add_paragraph()
r = p.add_run('Before SMOTE:  ')
r.bold = True; r.font.size = Pt(11)
r2 = p.add_run('Class 0 = 91,500  |  Class 1 = 8,500')
r2.font.size = Pt(11); r2.font.color.rgb = rgb(220,38,38)

p = doc.add_paragraph()
r = p.add_run('After SMOTE:   ')
r.bold = True; r.font.size = Pt(11)
r2 = p.add_run('Class 0 = 91,500  |  Class 1 = 91,500')
r2.font.size = Pt(11); r2.font.color.rgb = rgb(22,163,74)

doc.add_paragraph()

heading('3.3  Train / Test Split', level=2)
para('80% training  |  20% testing  |  random_state=42  |  stratified split', size=11, space_after=8)

heading('3.4  Feature Scaling', level=2)
para(
    'StandardScaler (zero mean, unit variance) is applied to models sensitive to '
    'feature magnitude: Logistic Regression, KNN, AdaBoost, and ANN. '
    'Tree-based models (Decision Tree, Random Forest, XGBoost) use unscaled data.',
    size=11
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. MODELS
# ═══════════════════════════════════════════════════════════════════════════════
heading('4. Models Used', level=1)
hline()

models_info = [
    ('1', 'Logistic Regression',
     'Linear model that estimates class probability using the sigmoid function. Serves as a strong baseline.',
     'max_iter=1000'),
    ('2', 'K-Nearest Neighbors (KNN)',
     'Instance-based learner that classifies a sample by majority vote of its k nearest neighbours in feature space.',
     'n_neighbors=5'),
    ('3', 'AdaBoost',
     'Boosting ensemble that sequentially trains weak learners (shallow Decision Trees), each correcting errors of the previous.',
     'n_estimators=50, base: DecisionTree(max_depth=1)'),
    ('4', 'Gaussian Naïve Bayes',
     'Probabilistic classifier based on Bayes\' theorem, assuming feature independence and Gaussian distributions.',
     'Default parameters'),
    ('5', 'Decision Tree',
     'Non-parametric model that splits data recursively based on feature thresholds to form a tree of decisions.',
     'max_depth=18'),
    ('6', 'Random Forest',
     'Bagging ensemble of many Decision Trees trained on random feature subsets. Reduces overfitting through averaging.',
     'n_estimators=200, max_depth=20'),
    ('7', 'XGBoost',
     'Gradient-boosted tree ensemble with regularisation. Highly optimised for speed and performance.',
     'n_estimators=200, lr=0.05, max_depth=4, subsample=0.8'),
    ('8', 'ANN (Neural Network)',
     'Three-layer feedforward network trained with Adam optimiser and binary cross-entropy loss for 100 epochs.',
     'Layers: Dense(32,relu) → Dense(16,relu) → Dense(1,sigmoid)'),
    ('9', 'Ensemble (RF + XGB + ANN)',
     'Soft-voting ensemble that averages the predicted probabilities of Random Forest, XGBoost, and ANN.',
     'Simple average of three model probabilities'),
]

for num, name, desc, params in models_info:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{num}.  {name}')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = rgb(79,70,229)
    para(desc, size=11, space_before=0, space_after=2)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    r_lbl = p2.add_run('Parameters: ')
    r_lbl.bold = True; r_lbl.font.size = Pt(10); r_lbl.font.color.rgb = rgb(100,116,139)
    r_val = p2.add_run(params)
    r_val.font.size = Pt(10); r_val.font.color.rgb = rgb(100,116,139)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
heading('5. Model Results & Comparison', level=1)
hline()
para(
    'All models were evaluated on the same held-out 20% test set. Metrics reported: '
    'Accuracy, Precision, Recall, F1-Score, and ROC-AUC.',
    size=11, space_after=10
)

result_cols = ['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'ROC-AUC']
result_rows = [
    ['Logistic Regression',   '91.46%', '0.9153', '0.9138', '0.9145', '0.9766'],
    ['K-Nearest Neighbors',   '94.59%', '0.9329', '0.9610', '0.9467', '0.9831'],
    ['AdaBoost',              '92.59%', '0.9136', '0.9409', '0.9270', '0.9855'],
    ['Naive Bayes',           '83.55%', '0.7741', '0.9474', '0.8520', '0.9305'],
    ['Decision Tree',         '95.57%', '0.9585', '0.9526', '0.9555', '0.9831'],
    ['Random Forest',         '96.67%', '0.9670', '0.9663', '0.9667', '0.9961'],
    ['XGBoost',               '96.53%', '0.9819', '0.9481', '0.9647', '0.9955'],
    ['ANN',                   '94.39%', '—',      '—',      '—',      '—'],
    ['Ensemble (RF+XGB+ANN)', '97.67%', '0.9825', '0.9707', '0.9766', '0.9975'],
]
highlight_row = 8   # Ensemble

t = doc.add_table(rows=1, cols=6)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(result_cols):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, '1E40AF')
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(h)
    run.bold = True; run.font.color.rgb = rgb(255,255,255); run.font.size = Pt(10)

for i, row_data in enumerate(result_rows):
    row = t.add_row()
    is_best = (i == highlight_row)
    bg = '14532D' if is_best else ('EEF2FF' if i % 2 == 0 else 'FFFFFF')
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p2.add_run(val)
        run.font.size = Pt(10)
        if is_best:
            run.bold = True
            run.font.color.rgb = rgb(74, 222, 128)
        elif j == 0:
            run.bold = True
            run.font.color.rgb = rgb(79,70,229)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('Best Model: ')
r1.bold = True; r1.font.size = Pt(11)
r2 = p.add_run('Ensemble (RF + XGB + ANN) — Accuracy 97.67%, ROC-AUC 0.9975')
r2.font.size = Pt(11); r2.font.color.rgb = rgb(22,163,74); r2.bold = True

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. ANN
# ═══════════════════════════════════════════════════════════════════════════════
heading('6. ANN Architecture & Training', level=1)
hline()

heading('6.1  Network Architecture', level=2)
ann_rows = [
    ['Input Layer',  '13 features (one per encoded column)',          '—'],
    ['Hidden Layer 1', '32 neurons, ReLU activation',                'Dense(32, relu)'],
    ['Hidden Layer 2', '16 neurons, ReLU activation',                'Dense(16, relu)'],
    ['Output Layer',   '1 neuron, Sigmoid activation (binary out)',   'Dense(1, sigmoid)'],
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(['Layer', 'Description', 'Keras Config']):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, '312E81')
    p2 = cell.paragraphs[0]
    run = p2.add_run(h)
    run.bold = True; run.font.color.rgb = rgb(255,255,255); run.font.size = Pt(10)
for i, row_data in enumerate(ann_rows):
    row = t.add_row()
    bg = 'F5F3FF' if i % 2 == 0 else 'FFFFFF'
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        p2 = cell.paragraphs[0]
        run = p2.add_run(val)
        run.font.size = Pt(10)
        if j == 0: run.bold = True

doc.add_paragraph()

heading('6.2  Training Configuration', level=2)
config = [
    ('Optimizer',        'Adam'),
    ('Loss Function',    'Binary Cross-Entropy'),
    ('Epochs',           '100'),
    ('Batch Size',       '16'),
    ('Validation Split', '20% of training data'),
    ('Best Val Accuracy','94.39%  (achieved at Epoch 98)'),
]
for lbl, val in config:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{lbl}: ')
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(val)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. ENSEMBLE
# ═══════════════════════════════════════════════════════════════════════════════
heading('7. Ensemble Model', level=1)
hline()
para(
    'The ensemble combines Random Forest, XGBoost, and ANN using soft voting — '
    'the predicted probabilities from all three models are averaged and a threshold '
    'of 0.5 is applied to produce the final binary prediction.',
    size=11, space_after=8
)
para('Formula:', bold=True, size=11, space_after=4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run('P_ensemble = ( P_RF + P_XGB + P_ANN ) / 3')
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = rgb(79,70,229)

ens_metrics = [
    ('Accuracy',  '97.67%'),
    ('Precision', '98.25%'),
    ('Recall',    '97.07%'),
    ('F1-Score',  '97.66%'),
    ('ROC-AUC',   '0.9975'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Metric', 'Score']):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, '14532D')
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(h)
    run.bold = True; run.font.color.rgb = rgb(255,255,255); run.font.size = Pt(11)
for lbl, val in ens_metrics:
    row = t.add_row()
    set_cell_bg(row.cells[0], 'F0FDF4')
    set_cell_bg(row.cells[1], 'DCFCE7')
    for j, v in enumerate([lbl, val]):
        p2 = row.cells[j].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(v)
        run.bold = True; run.font.size = Pt(11); run.font.color.rgb = rgb(22,163,74)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. WEB APPLICATION
# ═══════════════════════════════════════════════════════════════════════════════
heading('8. Web Application', level=1)
hline()
para(
    'Two Flask web applications provide interactive access to the project results '
    'and the prediction model.',
    size=11, space_after=10
)

heading('8.1  Dashboard  (Port 5000)', level=2)
para('URL: http://127.0.0.1:5000', bold=True, size=11, color=rgb(79,70,229))
bullet('Accuracy summary cards for all 9 models')
bullet('Interactive bar charts — Accuracy, F1-Score, ROC-AUC')
bullet('Full metrics comparison table')
bullet('ROC curve plots for every model')
bullet('ANN training accuracy and loss curves over 100 epochs')

doc.add_paragraph()

heading('8.2  Patient Prediction Form  (Port 5001)', level=2)
para('URL: http://127.0.0.1:5001', bold=True, size=11, color=rgb(79,70,229))
bullet('Clinical input form: Gender, Age, BMI, HbA1c, Blood Glucose, Hypertension, Heart Disease, Smoking')
bullet('Real-time prediction using the saved Random Forest model (96.67% accuracy)')
bullet('Displays risk probability bar (0 – 100%)')
bullet('Color-coded result: green (low risk) / red (high risk)')
bullet('Personalised health recommendations based on the prediction')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading('9. Project File Structure', level=1)
hline()

files = [
    ('diabetes_prediction_dataset.csv', 'Raw dataset — 100,000 patient records'),
    ('run_diabetes.py',                 'Full training script — trains all 9 models and saves plots'),
    ('save_models.py',                  'Trains and saves the Random Forest model + scaler to disk'),
    ('app.py',                          'Flask dashboard web app (port 5000)'),
    ('predict_app.py',                  'Flask patient prediction web app (port 5001)'),
    ('start.py',                        'One-click launcher — starts both Flask apps simultaneously'),
    ('plots/',                          'All generated plots (ROC curves, ANN accuracy/loss, ensemble ROC)'),
    ('saved_model/rf_model.pkl',        'Trained Random Forest model (saved with joblib)'),
    ('saved_model/scaler.pkl',          'Fitted StandardScaler (saved with joblib)'),
    ('saved_model/feature_cols.pkl',    'List of encoded feature column names'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['File / Folder', 'Purpose']):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, '0F172A')
    p2 = cell.paragraphs[0]
    run = p2.add_run(h)
    run.bold = True; run.font.color.rgb = rgb(255,255,255); run.font.size = Pt(10)
for i, (fname, desc) in enumerate(files):
    row = t.add_row()
    bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    r1 = row.cells[0].paragraphs[0].add_run(fname)
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = rgb(79,70,229)
    r2 = row.cells[1].paragraphs[0].add_run(desc)
    r2.font.size = Pt(9.5)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. HOW TO RUN
# ═══════════════════════════════════════════════════════════════════════════════
heading('10. How to Run the Project', level=1)
hline()

heading('10.1  Prerequisites', level=2)
for pkg in ['pandas', 'numpy', 'scikit-learn', 'imbalanced-learn', 'xgboost',
            'tensorflow', 'matplotlib', 'seaborn', 'flask', 'joblib']:
    bullet(pkg)

doc.add_paragraph()

heading('10.2  First-Time Setup', level=2)
steps = [
    ('Step 1', 'Train all models and generate plots',  'python run_diabetes.py'),
    ('Step 2', 'Save the Random Forest model to disk', 'python save_models.py'),
    ('Step 3', 'Launch both web apps',                 'python start.py'),
]
for label, desc, cmd in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{label}: ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = rgb(79,70,229)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_after = Pt(8)
    r3 = p2.add_run(f'> {cmd}')
    r3.font.size = Pt(10); r3.font.color.rgb = rgb(22,163,74)

heading('10.3  Subsequent Runs (models already saved)', level=2)
para('Simply run:', size=11, space_after=4)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r = p.add_run('> python start.py')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = rgb(22,163,74)
para('Both browser tabs open automatically.', size=11, space_before=4)

doc.add_paragraph()

heading('10.4  Accessing the Applications', level=2)
apps = [
    ('Dashboard',      'http://127.0.0.1:5000', 'All model metrics, charts, and ROC curves'),
    ('Patient Form',   'http://127.0.0.1:5001', 'Enter patient data and get diabetes prediction'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(['Application', 'URL', 'Purpose']):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, '1E3A8A')
    p2 = cell.paragraphs[0]
    run = p2.add_run(h)
    run.bold = True; run.font.color.rgb = rgb(255,255,255); run.font.size = Pt(10)
for i, (app_name, url, purpose) in enumerate(apps):
    row = t.add_row()
    bg = 'EFF6FF' if i == 0 else 'F0FDF4'
    for j, val in enumerate([app_name, url, purpose]):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        p2 = cell.paragraphs[0]
        run = p2.add_run(val)
        run.font.size = Pt(10)
        if j == 0: run.bold = True
        if j == 1: run.font.color.rgb = rgb(79,70,229)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading('11. Conclusion', level=1)
hline()
para(
    'This project successfully demonstrates the power of ensemble learning for medical '
    'classification tasks. By combining three complementary models — Random Forest, '
    'XGBoost, and an Artificial Neural Network — the ensemble achieves a 97.67% accuracy '
    'and a ROC-AUC of 0.9975, outperforming every individual model.',
    size=11, space_after=8
)
para('Key takeaways:', bold=True, size=11, space_after=4)
bullet('SMOTE is critical for handling class imbalance in medical datasets')
bullet('Ensemble methods consistently outperform individual classifiers')
bullet('Random Forest and XGBoost are the strongest individual models (>96% accuracy)')
bullet('The ANN requires more data or regularisation to match tree-based models')
bullet('Naive Bayes has the lowest accuracy (83.55%) but the highest recall — useful when minimising false negatives is critical')
bullet('The web application makes the model accessible to non-technical users')

doc.add_paragraph()
para(
    'The project is fully reproducible: all scripts, saved models, and web applications '
    'are included. Future improvements could include hyperparameter tuning (GridSearchCV), '
    'additional features (family history, physical activity), SHAP explainability, '
    'and deployment to a public cloud platform.',
    size=11
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = 'D:/sugarrr/Diabetes_Prediction_Project_Report.docx'
doc.save(out)
print(f"Document saved: {out}")
